"""
文档解析器

支持 PDF、Word (.docx)、TXT、Markdown 格式的文本提取。
提供两种解析模式：
- parse(): 返回纯文本（向后兼容）
- parse_structured(): 返回带位置信息的结构化结果
"""
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


@dataclass
class TextBlock:
    """
    解析后的文本块（含位置信息）

    Attributes:
        content: 文本内容
        locationType: 位置类型 ("page" | "line" | "section" | "paragraph")
        locationValue: 位置值 ("3" / "15" / "第三章 假期规定")
    """
    content: str
    locationType: str = ""
    locationValue: str = ""


@dataclass
class ParseResult:
    """
    文档解析结果

    Attributes:
        blocks: 带位置信息的文本块列表
        fullText: 完整纯文本（向后兼容）
    """
    blocks: List[TextBlock] = field(default_factory=list)
    fullText: str = ""


class DocumentParser:
    """
    文档解析器

    根据文件扩展名自动选择解析策略。
    支持两种输出模式：纯文本和结构化位置信息。
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    @staticmethod
    def parse(file_path: str) -> str:
        """
        解析文档并提取纯文本（向后兼容）

        Args:
            file_path: 文件路径

        Returns:
            提取的纯文本内容
        """
        result = DocumentParser.parse_structured(file_path)
        return result.fullText

    @staticmethod
    def parse_structured(file_path: str) -> ParseResult:
        """
        解析文档并返回带位置信息的结构化结果

        Args:
            file_path: 文件路径

        Returns:
            ParseResult 包含 TextBlock 列表和完整文本

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in DocumentParser.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}，支持: {DocumentParser.SUPPORTED_EXTENSIONS}")

        parsers = {
            ".pdf": DocumentParser._parse_pdf_structured,
            ".docx": DocumentParser._parse_docx_structured,
            ".txt": DocumentParser._parse_text_structured,
            ".md": DocumentParser._parse_markdown_structured,
        }
        parser = parsers.get(ext)
        if not parser:
            raise ValueError(f"未实现 {ext} 格式的解析器")

        logger.info(f"解析文档: {path.name} ({ext})")
        return parser(file_path)

    @staticmethod
    def _parse_pdf_structured(file_path: str) -> ParseResult:
        """
        解析 PDF 文件（结构化）

        每页作为一个 TextBlock，locationType="page"

        Args:
            file_path: PDF 文件路径

        Returns:
            ParseResult
        """
        import pdfplumber

        blocks: List[TextBlock] = []
        pages_text = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    page_num = str(i + 1)
                    blocks.append(TextBlock(
                        content=text.strip(),
                        locationType="page",
                        locationValue=page_num,
                    ))
                    pages_text.append(f"[第{page_num}页]\n{text.strip()}")

        return ParseResult(
            blocks=blocks,
            fullText="\n\n".join(pages_text),
        )

    @staticmethod
    def _parse_docx_structured(file_path: str) -> ParseResult:
        """
        解析 Word (.docx) 文件（结构化）

        每个段落作为一个 TextBlock：
        - 标题段落: locationType="section", locationValue=标题文本
        - 正文段落: locationType="paragraph", locationValue=段落序号

        Args:
            file_path: Word 文件路径

        Returns:
            ParseResult
        """
        from docx import Document

        doc = Document(file_path)
        blocks: List[TextBlock] = []
        paragraphs_text = []
        para_index = 0
        current_section = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            para_index += 1
            style_name = para.style.name if para.style else ""

            if style_name.startswith("Heading"):
                # 标题段落
                current_section = text
                blocks.append(TextBlock(
                    content=f"## {text}",
                    locationType="section",
                    locationValue=text,
                ))
                paragraphs_text.append(f"## {text}")
            else:
                # 正文段落
                location = current_section if current_section else f"第{para_index}段"
                blocks.append(TextBlock(
                    content=text,
                    locationType="section" if current_section else "paragraph",
                    locationValue=location,
                ))
                paragraphs_text.append(text)

        # 表格内容
        table_index = 0
        for table in doc.tables:
            table_index += 1
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip("| "):
                    blocks.append(TextBlock(
                        content=row_text,
                        locationType="paragraph",
                        locationValue=f"表格{table_index}",
                    ))
                    paragraphs_text.append(row_text)

        return ParseResult(
            blocks=blocks,
            fullText="\n\n".join(paragraphs_text),
        )

    @staticmethod
    def _read_text_safe(file_path: str) -> str:
        """
        安全读取文本文件，尝试多种编码

        Args:
            file_path: 文件路径

        Returns:
            文本内容
        """
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        # 最终回退：忽略无法解码的字符
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    @staticmethod
    def _parse_text_structured(file_path: str) -> ParseResult:
        """
        解析纯文本文件（结构化）

        按段落切分，每个 TextBlock 记录起始行号。
        locationType="line", locationValue="起始行号"

        Args:
            file_path: TXT 文件路径

        Returns:
            ParseResult
        """
        lines = DocumentParser._read_text_safe(file_path).splitlines(keepends=True)
        blocks: List[TextBlock] = []
        block_texts = []
        current_lines = []
        start_line = 1

        for i, line in enumerate(lines):
            stripped = line.rstrip("\n").rstrip("\r")
            current_lines.append(stripped)
            # 空行作为段落分隔
            if stripped.strip() == "":
                if len(current_lines) > 1:
                    # 前面有非空行，保存当前段落
                    content = "\n".join(current_lines[:-1])
                    if content.strip():
                        blocks.append(TextBlock(
                            content=content,
                            locationType="line",
                            locationValue=str(start_line),
                        ))
                        block_texts.append(content)
                    start_line = i + 2
                current_lines = []
                start_line = i + 2

        # 处理最后一段
        if current_lines:
            content = "\n".join(current_lines)
            if content.strip():
                blocks.append(TextBlock(
                    content=content,
                    locationType="line",
                    locationValue=str(start_line),
                ))
                block_texts.append(content)

        return ParseResult(
            blocks=blocks,
            fullText="\n\n".join(block_texts),
        )

    @staticmethod
    def _parse_markdown_structured(file_path: str) -> ParseResult:
        """
        解析 Markdown 文件（结构化）

        同 TXT 处理逻辑，按段落切分，记录起始行号。

        Args:
            file_path: Markdown 文件路径

        Returns:
            ParseResult
        """
        # Markdown 按行号追踪，逻辑与 TXT 一致
        return DocumentParser._parse_text_structured(file_path)
